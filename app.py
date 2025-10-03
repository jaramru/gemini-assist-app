import os
import re
from io import BytesIO
from datetime import datetime

import streamlit as st
import pandas as pd
import google.generativeai as genai

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


# ================== Config página ==================
st.set_page_config(page_title="Gemini Assist – Informe Predictivo", layout="wide")

# Cabecera con logo
try:
    st.image("images/logo.png", width=150)
except Exception:
    pass

st.title("🔧 Gemini Assist – Informe Predictivo de Mantenimiento (modo chat)")


# ================== API KEY helpers ==================
def get_api_key() -> str | None:
    """Lee la API key de st.secrets o variables de entorno (nombres comunes)."""
    try:
        for k in ("GOOGLE_API_KEY", "API_KEY", "GEMINI_API_KEY"):
            if k in st.secrets and str(st.secrets[k]).strip():
                return str(st.secrets[k]).strip()
    except Exception:
        pass
    for k in ("GOOGLE_API_KEY", "API_KEY", "GEMINI_API_KEY"):
        v = os.getenv(k, "").strip()
        if v:
            return v
    return None


# ================== Limpieza / formato ==================
_bullet_regex = re.compile(r"^\s*[-•]\s*")

def normaliza_numeracion(linea: str) -> str:
    # '1. 1. Título' -> '1. Título'
    return re.sub(r"^(\s*\d+\.\s+)(\d+\.\s+)+", r"\1", linea)

def limpiar_texto_base(texto: str) -> str:
    """Quita **negritas** markdown, asteriscos, homogeneiza viñetas a '• ' y corrige guiones raros."""
    texto = re.sub(r"\*\*(.*?)\*\*", r"\1", texto)
    lineas = []
    for raw in texto.splitlines():
        l = raw.rstrip()
        l = normaliza_numeracion(l)
        if re.match(r"^\s*[\*\-]\s+", l):
            l = re.sub(r"^\s*[\*\-]\s+", "• ", l)
        if "*" in l:
            l = l.replace("*", "")
        l = l.replace("–", "-").replace("—", "-")
        lineas.append(l)
    return "\n".join(lineas).strip()

def es_encabezado(linea: str) -> bool:
    if re.match(r"^\d+\.\s", linea.strip()):
        return True
    patrones = [
        "Acciones Preventivas", "Acciones preventivas",
        "Estimación de ahorro", "Estimacion de ahorro",
        "Panel de alertas", "Informe ejecutivo",
    ]
    return any(pat.lower() in linea.lower() for pat in patrones)


# ================== DOCX (portada + contenido) ==================
def generar_word(informe: str) -> BytesIO:
    doc = Document()

    sec = doc.sections[0]
    sec.top_margin = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)

    # Portada
    try:
        p_logo = doc.add_paragraph()
        r_logo = p_logo.add_run()
        r_logo.add_picture("images/logo.png", width=Inches(2))
        p_logo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    except Exception:
        pass

    tit = doc.add_paragraph("Gemini Assist")
    tit.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    rt = tit.runs[0]; rt.bold = True; rt.font.size = Pt(26)

    subt = doc.add_paragraph("Informe Predictivo de Mantenimiento Hospitalario")
    subt.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    rs = subt.runs[0]; rs.font.size = Pt(13); rs.font.color.rgb = RGBColor(90, 90, 90)

    fecha = doc.add_paragraph(datetime.now().strftime("%d/%m/%Y"))
    fecha.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    rf = fecha.runs[0]; rf.font.size = Pt(11); rf.font.color.rgb = RGBColor(90, 90, 90)

    doc.add_page_break()

    # Contenido
    texto = limpiar_texto_base(informe)
    for raw in texto.splitlines():
        l = raw.strip()
        if not l:
            continue

        if es_encabezado(l):
            p = doc.add_paragraph()
            r = p.add_run(l)
            r.bold = True
            r.font.size = Pt(13)
            p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            continue

        if _bullet_regex.match(l):
            p = doc.add_paragraph(l[_bullet_regex.match(l).end():], style="List Bullet")
            p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            for run in p.runs:
                run.font.size = Pt(11)
            continue

        p = doc.add_paragraph(l)
        p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        for run in p.runs:
            run.font.size = Pt(11)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ================== System Instructions ==================
SYSTEM_INSTRUCTIONS = """
Eres Gemini Assist, un sistema experto en mantenimiento hospitalario.

Reglas de estilo:
- Estilo neutro, profesional y en blanco y negro.
- No uses asteriscos (*) ni emojis.
- Usa títulos claros (ej. '1. Acciones preventivas...') y listas con viñetas (•) cuando corresponda.
- Redacción clara y concisa, justificada.

Estructura obligatoria del informe (en este orden):
1) Acciones preventivas para los 3 activos más críticos.
2) Estimación de ahorro en € y horas si se aplican esas medidas.
3) Panel de alertas clasificando cada activo en: Bajo, Medio o Alto.
4) Informe ejecutivo final (máximo 5 líneas).

Notas:
- No incluyas Markdown decorativo (##, ###) ni negritas con ** **.
- Evita símbolos raros (✔, ❌, etc.).
- No repitas numeración (nada de '1. 1.').
"""


# ================== Resolver modelo disponible ==================
PREFERRED_MODELS = [
    "gemini-2.5-flash",
    "gemini-2.5-pro",
    "gemini-2.0-flash",
    "gemini-2.0-flash-lite",
    "gemini-1.5-flash",
    "gemini-1.5-pro",
    # Fallbacks para SDKs viejos (v1beta):
    "gemini-1.0-pro",
    "gemini-pro",
]

def resolve_model_id():
    """Intenta listar modelos y elegir el mejor disponible para generateContent."""
    try:
        names = []
        for m in genai.list_models():
            n = getattr(m, "name", "")
            if n.startswith("models/"):
                n = n.split("/", 1)[1]
            methods = set(getattr(m, "supported_generation_methods", []) or [])
            if "generateContent" in methods or not methods:
                if n:
                    names.append(n)
        for wanted in PREFERRED_MODELS:
            if wanted in names:
                return wanted
    except Exception:
        pass
    return "gemini-1.0-pro"


def crear_modelo_con_fallback(model_id: str):
    """
    Crea el modelo intentando usar system_instruction.
    Si el SDK es antiguo y no lo soporta, hace fallback sin ese argumento.
    """
    try:
        return genai.GenerativeModel(
            model_name=model_id,
            system_instruction=SYSTEM_INSTRUCTIONS
        ), True
    except TypeError:
        return genai.GenerativeModel(model_name=model_id), False


# ================== UI: subir Excel ==================
st.subheader("📎 Sube el archivo de activos (Excel)")
uploaded_file = st.file_uploader("Arrastra y suelta, o pulsa en **Browse files**", type=["xlsx"])

if uploaded_file:
    try:
        df = pd.read_excel(uploaded_file)
        st.success("✅ Archivo cargado correctamente")
        st.dataframe(df.head(50), use_container_width=True)

        # Guardamos la tabla en sesión una sola vez (contexto del chat)
        if "tabla_texto" not in st.session_state:
            st.session_state.tabla_texto = df.to_string(index=False)

    except Exception as e:
        st.error(f"❌ No se pudo leer el Excel: {e}")
        st.stop()
else:
    st.info("Carga un archivo .xlsx para comenzar.")
    st.stop()


# ================== Configurar modelo + chat ==================
api_key = get_api_key()
if not api_key:
    st.error(
        "❌ No se encontró la API KEY. Configúrala en Streamlit Cloud → Settings → Secrets con:\n\n"
        'GOOGLE_API_KEY="tu_clave"  (o API_KEY)'
    )
    st.stop()

try:
    genai.configure(api_key=api_key)
except Exception as e:
    st.error(f"⚠️ Error configurando la API KEY: {e}")
    st.stop()

MODEL_ID = resolve_model_id()
model, tiene_system = crear_modelo_con_fallback(MODEL_ID)

# Inicializamos chat en sesión (persistente) — sin mostrar respuestas técnicas
if "chat" not in st.session_state:
    st.session_state.chat = model.start_chat(history=[])
    # Enviamos la tabla como contexto inicial (silencioso)
    primer_prompt = f"""
Analiza y guarda en memoria la siguiente tabla de activos hospitalarios (texto). La usaremos como contexto en esta conversación.

TABLA:
{st.session_state.tabla_texto}
"""
    try:
        if tiene_system:
            st.session_state.chat.send_message(primer_prompt)
        else:
            st.session_state.chat.send_message(SYSTEM_INSTRUCTIONS + "\n\n" + primer_prompt)
    except Exception as e:
        st.error(f"⚠️ No se pudo inicializar el contexto: {e}")


# ================== UI de interacción minimalista ==================
user_msg = st.text_area(
    label="",  # sin etiqueta visible
    placeholder="Escribe aquí tu instrucción o ajuste para el informe…",
    height=140,
    label_visibility="collapsed"
)

col1, col2, col3 = st.columns(3)
with col1:
    enviar = st.button("➡️ Enviar", type="primary")
with col2:
    generar_final = st.button("🧾 Generar Informe Final (Word)")
with col3:
    reset_chat = st.button("🧹 Reiniciar")

if reset_chat:
    st.session_state.chat = model.start_chat(history=[])
    try:
        primer_prompt = f"""
Analiza y guarda en memoria la siguiente tabla de activos hospitalarios (texto). La usaremos como contexto en esta conversación.

TABLA:
{st.session_state.tabla_texto}
"""
        if tiene_system:
            st.session_state.chat.send_message(primer_prompt)
        else:
            st.session_state.chat.send_message(SYSTEM_INSTRUCTIONS + "\n\n" + primer_prompt)
        st.success("Conversación reiniciada.")
    except Exception as e:
        st.error(f"⚠️ No se pudo reiniciar el contexto: {e}")

if enviar and user_msg.strip():
    with st.spinner("Pensando..."):
        try:
            resp = st.session_state.chat.send_message(user_msg)
            respuesta = getattr(resp, "text", "") or ""
            respuesta_limpia = limpiar_texto_base(respuesta)

            # Mostrar respuesta con títulos en negrita y viñetas limpias
            vista = []
            for raw in respuesta_limpia.splitlines():
                l = raw.strip()
                if not l:
                    continue
                if es_encabezado(l):
                    vista.append(f"**{l}**")
                elif _bullet_regex.match(l):
                    vista.append(f"- {l[_bullet_regex.match(l).end():]}")
                else:
                    vista.append(l)
            st.markdown("\n\n".join(vista))
        except Exception as e:
            st.error(f"❌ Error en la respuesta del asistente: {e}")

# ================== Generar informe final (Word) ==================
if generar_final:
    with st.spinner("Generando documento final..."):
        try:
            pedido_final = (
                "Genera el informe final siguiendo las instrucciones del sistema "
                "(sin markdown ni emojis; títulos claros; viñetas con '•'; redacción justificada)."
            )
            resp_final = st.session_state.chat.send_message(pedido_final) if tiene_system \
                         else st.session_state.chat.send_message(SYSTEM_INSTRUCTIONS + "\n\n" + pedido_final)
            informe_final = getattr(resp_final, "text", "") or ""
            informe_final_limpio = limpiar_texto_base(informe_final)

            word_bytes = generar_word(informe_final_limpio)
            st.download_button(
                "⬇️ Descargar Informe (Word)",
                data=word_bytes,
                file_name="informe_predictivo.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        except Exception as e:
            st.error(f"❌ Error al generar el informe final: {e}")
